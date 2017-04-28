from pandas import DataFrame
import pandas as pd
import numpy as np
import math
from . import models
from django.shortcuts import render
from django.http import HttpResponse,HttpResponseRedirect,StreamingHttpResponse
import os
import json
import datetime
import csv
from decimal import *
from docx import Document
from docx.shared import Inches
#from . import hashuang

#多条件筛选：可自由选择要进行筛选的条件
def multi_analy(request):
	print("multi_analy")
	bookno=request.POST.get("bookno").upper();
	gk_no=request.POST.get("gk_no");
	SPECIFICATION_ask=str(request.POST.get("SPECIFICATION"));#钢种的筛选要求，native/all
	OPERATECREW_ask=str(request.POST.get("OPERATECREW"));#班别的筛选要求native/all/A/B/C
	station=request.POST.get("station");
	if gk_no !='blank':
		sentence_gk_no= " and gk_no='"+gk_no+"'"
	else:
		sentence_gk_no=''
	if OPERATESHIFT !='blank':
		sentence_OPERATESHIFT=" and OPERATESHIFT='"+OPERATESHIFT+"'"
	else:
		sentence_OPERATESHIFT=''
	if OPERATECREW !='blank':
		sentence_OPERATECREW=" and OPERATECREW='"+OPERATECREW+"'"
	else:
		sentence_OPERATECREW=''
	if station !='blank':
		sentence_station=" and station='"+station+"'"
	else:
		sentence_station=''
	sentence="SELECT HEAT_NO,"+bookno+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS WHERE HEAT_NO>'1500000'"+sentence_gk_no+sentence_OPERATESHIFT+sentence_OPERATECREW+sentence_station
	#print(sentence)
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]=sentence
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	print(scrapy_records[:5])
	contentVO={
		'title':'测试',
		'state':'success',
	}
	ana_result,ana_describe=num_describe(scrapy_records,bookno)
	contentVO['result']=ana_result
	contentVO['describe']=ana_describe
	return HttpResponse(json.dumps(contentVO),content_type='application/json')

#去括号取左侧的数（计算概率分布时需要用到）
def getA(s):
    s0 = list(s)
    data = []
    for a in s0:
        a1 = a.split(',')
        if len(a1)==1:
            data.append(a1[0][1:])
        else:
            data.append(a1[0][1:])
    return data

#去括号取右侧的数（计算概率分布时需要用到）    
def getB(s):
    s0 = list(s)
    data = []
    for a in s0:
        a1 = a.split(',')
        if len(a1)==1:
            data.append(a1[1][:-1])
        else:
            data.append(a1[1][:-1])
    return data	

#拼接区间（计算概率分布时需要用到） 
def union_section(section_point,sections):
    for i in range(len(section_point)):
        section=None
        if i<len(section_point)-1:
            section='('+str(section_point[i])+','+str(section_point[i+1])+']'
            sections.append(section)
    return sections 

#计算转炉含量
def cost(request):
	print('你好')
	#print("cost_success")
	prime_cost=request.POST.get("prime_cost");
	SPECIFICATION=request.POST.get("SPECIFICATION");#钢种
	OPERATECREW=request.POST.get("OPERATECREW");#班别
	time1=request.POST.get("time1");
	time2=request.POST.get("time2");

	xaxis=['C含量','SI含量','MN含量','P含量','S含量','重量','温度']
	xasis_fieldname=['C','SI','MN','P','S','STEELWGT','FINAL_TEMP_VALUE']

	#先查出当前炉次实际所属钢种及班别
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,SPECIFICATION,OPERATECREW FROM qg_user.PRO_BOF_HIS_ALLFIELDS WHERE HEAT_NO='"+prime_cost+"'";
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	if len(scrapy_records)==0:#无该炉次号
			contentVO['state']='error'
			return HttpResponse(json.dumps(contentVO),content_type='application/json')

	actual_SPECIFICATION=scrapy_records[0].get('SPECIFICATION',None)
	actual_OPERATECREW=scrapy_records[0].get('OPERATECREW',None)
	
	#确认实际筛选语句		
	str_select=""
	if SPECIFICATION=='native' and actual_SPECIFICATION !=None:#本钢种且本钢种不为空
		str_select=str_select+" and SPECIFICATION = '"+str(actual_SPECIFICATION)+"'"

	if OPERATECREW=='native' and actual_OPERATECREW !=None:
		str_select=str_select+" and OPERATECREW = '"+str(actual_OPERATECREW)+"'"
	elif OPERATECREW=='all':
		pass
	else:
		str_select=str_select+" and OPERATECREW = '"+OPERATECREW_ask+"'"

	#时间范围
	if time1 != '' and time2!='':
		str_select=str_select+"and to_char(MSG_DATE_PLAN,'yyyy-mm-dd')>'"+time1+"'and to_char(MSG_DATE_PLAN,'yyyy-mm-dd')<'"+time2+"'"
	else:
		pass
	print('str_select',str_select)



	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,nvl(C,0) as C,nvl(SI,0) as SI ,nvl(MN,0) as MN,nvl(P,0) as P ,nvl(S,0) as S, nvl(STEELWGT,0)as STEELWGT,nvl(FINAL_TEMP_VALUE,0)as FINAL_TEMP_VALUE FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+prime_cost+"'";
	# sqlVO["sql"]=sentence
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	contentVO={
		'title':'测试',
		'state':'success'
	}


	print(scrapy_records)

	for i in range(len(xasis_fieldname)):
		value = scrapy_records[0].get(xasis_fieldname[i],None)
		if value != None :
			scrapy_records[0][xasis_fieldname[i]] = float(value)
	frame=DataFrame(scrapy_records)

	
	yaxis=[frame.C[0],frame.SI[0],frame.MN[0],frame.P[0],frame.S[0],frame.STEELWGT[0],frame.FINAL_TEMP_VALUE[0]]
	print('xasis_fieldname',xasis_fieldname)#分析字段英文名
	print('yaxis',yaxis)#分析字段实际值
	offset_result=offset(xasis_fieldname,yaxis,str_select)#计算偏离程度函数的返回值
	#offset_result=offset(xasis_fieldname,yaxis)#计算偏离程度函数的返回值

	qualitative_offset_result=qualitative_offset(offset_result)#对偏离程度进行定性判断
	offset_resultlist=["%.2f%%"%(n*100) for n in list(offset_result)]
	print('offset_result',offset_result)#分析字段偏离程度
	print('qualitative_offset_result',qualitative_offset_result)#偏离程度进行定性判断
	print('offset_resultlist',offset_resultlist)#分析字段偏离程度格式化

	j=0;
	for i in range(len(offset_result)):
		if abs(offset_result[i])<0.2:
			j+=1
	if j==7:
		contentVO['normal']='true'
	else:
		contentVO['normal']='flase'	

	ana_result={}
	ana_result['heat_no']=frame.HEAT_NO[0]#炉次号
	ana_result['xname']=xaxis#字段中文名字
	ana_result['xEnglishname']=xasis_fieldname#字段英文名字
	ana_result['yvalue']=yaxis#该炉次字段的实际值
	ana_result['str_select']=str_select#筛选条件
	ana_result['attribute']='含量'
	ana_result['offset_result']=offset_resultlist#各字段的偏离程度值
	ana_result['qualitative_offset_result']=qualitative_offset_result#各字段的偏离程度定性判断结果
	contentVO['result']=ana_result
	return HttpResponse(json.dumps(contentVO),content_type='application/json')

def offset(xasis_fieldname,yaxis,str_select):
	offset_result=[]
	sqlVO={}
	sqlVO["db_name"]="l2own"
	print('len(xasis_fieldname)',len(xasis_fieldname))

	parameters=["MAX_VALUE","MIN_VALUE","DESIRED_VALUE","STANDARD_DEVIATION",'NUMERICAL_LOWER_BOUND','NUMERICAL_UPPER_BOUND','IF_FIVENUMBERSUMMARY']
	for i in range (len(xasis_fieldname)):#实际计算范围为从0到len(xasis_fieldname)-1
		# if yaxis[i]==0 or yaxis[i] is None:
		# 	offset_result.append(None)
		# 	continue;
		sqlVO["sql"]="SELECT MAX_VALUE,MIN_VALUE,DESIRED_VALUE,STANDARD_DEVIATION,NUMERICAL_LOWER_BOUND,NUMERICAL_UPPER_BOUND,IF_FIVENUMBERSUMMARY FROM qg_user.PRO_BOF_HIS_ALLSTRUCTURE where DATA_ITEM_EN = \'"+xasis_fieldname[i]+"\'"
		print(sqlVO["sql"])
		scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
		print(scrapy_records)
		# print(temp_array)
		# print(yaxis[i])
		# print(isinstance(yaxis[i],float))#判断数据类型
		for j in range (len(parameters)):
			value = scrapy_records[0].get(parameters[j],None)
			if value != None and value != 'null':
				scrapy_records[0][parameters[j]] = float(value)


		if str_select=='':#相当于无筛选条件,直接根据最大最小值计算偏离程度
			#如果实际值大于最大值，则按最大值计算，偏离程度为50% ; 如果实际值小于最小值，则按最小值处理。偏离程度为-50%
			if float(yaxis[i])>scrapy_records[0]['MAX_VALUE']:
				adjusted_yaxis=scrapy_records[0]['MAX_VALUE']
			elif float(yaxis[i])<scrapy_records[0]['MIN_VALUE']:
				adjusted_yaxis=scrapy_records[0]['MIN_VALUE']
			else :
				adjusted_yaxis=float(yaxis[i])


			try:
				# temp_value=(float(yaxis[i])-scrapy_records[0]['DESIRED_VALUE'])/(scrapy_records[0]['MAX_VALUE']-scrapy_records[0]['MIN_VALUE'])
				temp_value=(adjusted_yaxis-scrapy_records[0]['MIN_VALUE'])/(scrapy_records[0]['MAX_VALUE']-scrapy_records[0]['MIN_VALUE'])-0.5
			except:
				temp_value=None
			offset_result.append(temp_value)

		else:#有筛选条件,需要动态经过上下限、五数等清洗后计算最大最小值
			#bound_select:上下限筛选条件
			bound_select =' and '+xasis_fieldname[i]+'>='+str(scrapy_records[0]['NUMERICAL_LOWER_BOUND']) +' and '+xasis_fieldname[i]+'<='+str(scrapy_records[0]['NUMERICAL_UPPER_BOUND'])
			sqlVO1={}
			sqlVO1["db_name"]="l2own"
			sqlVO1["sql"]="SELECT HEAT_NO,"+xasis_fieldname[i]+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no>'1500000'"+str_select+bound_select
			print(sqlVO1["sql"])
			scrapy_records1=models.BaseManage().direct_select_query_sqlVO(sqlVO1)

			if len(scrapy_records1)<10:#对比数据条数不足10条，表示无法进行偏离程度计算
				offset_result.append(None)
				continue

			for k in range(len(scrapy_records1)):
					value = scrapy_records1[k].get(xasis_fieldname[i],None)
					if value != None :
						scrapy_records1[k][xasis_fieldname[i]] = float(value)
			frame=DataFrame(scrapy_records1)
			df=frame.sort_values(by=xasis_fieldname[i])
			print(df)
			#进行五数清洗

			if scrapy_records[0]['IF_FIVENUMBERSUMMARY']==1:
				if len(scrapy_records1)>150:				
					clean=Wushu(df[xasis_fieldname[i]])['result']
				else:
					clean=df[xasis_fieldname[i]]	
			else:
				clean=df[xasis_fieldname[i]]
			print('五数清洗后的数据条数：',len(clean))
			if len(clean)<10:#数据不足10条
				offset_result.append(None)
				continue
			# print('clean结果',clean)

			# dataclean_result={}
			avg_value=np.mean(clean)
			print("期望值",avg_value)
			#标准差
			std_value=np.std(clean)
			print("标准差",std_value)
			#方差
			var_value=np.var(clean)
			print("方差",var_value)
			#normx,normy=Norm_dist(avg_value,var_value)
			print("min",clean.min())
			print("max",clean.max())

			# temp_array=data_clean(scrapy_records1,xasis_fieldname[i])#字段清洗后统计情况
			# print(temp_array)
			print(yaxis[i])
			# print(isinstance(yaxis[i],float))#判断数据类型

			#如果实际值大于最大值，则按最大值计算，偏离程度为50% ; 如果实际值小于最小值，则按最小值处理。偏离程度为-50%
			if float(yaxis[i])>clean.max():
				adjusted_yaxis=clean.max()
			elif float(yaxis[i])<clean.min():
				adjusted_yaxis=clean.min()
			else :
				adjusted_yaxis=float(yaxis[i])


			try:
				temp_value=(adjusted_yaxis-avg_value)/(clean.max()-clean.min())
			except:
				temp_value=None
			offset_result.append(temp_value)


	print('offset_result',offset_result)
	return offset_result

#通过从数据库中查询期望等参数图中的偏离程度
def offset_fu(xasis_fieldname,yaxis):
	offset_result=[]
	sqlVO={}
	sqlVO["db_name"]="l2own"
	print('len(xasis_fieldname)',len(xasis_fieldname))
	parameters=["MAX_VALUE","MIN_VALUE","DESIRED_VALUE","STANDARD_DEVIATION"]
	for i in range (0,len(xasis_fieldname)):#实际计算范围为从0到len(xasis_fieldname)-1
		sqlVO["sql"]="SELECT MAX_VALUE,MIN_VALUE,DESIRED_VALUE,STANDARD_DEVIATION FROM qg_user.PRO_BOF_HIS_ALLSTRUCTURE where DATA_ITEM_EN = \'"+xasis_fieldname[i]+"\'"
		print(sqlVO["sql"])
		scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
		print(scrapy_records)
		for j in range (4):
			value = scrapy_records[0].get(parameters[j],None)
			if value != None and value != 'null':
				scrapy_records[0][parameters[j]] = float(value)
		try:
			temp_value=((float(yaxis[i])-scrapy_records[0]['MIN_VALUE'])/(scrapy_records[0]['MAX_VALUE']-scrapy_records[0]['MIN_VALUE']))- 0.5
		except:
			temp_value=None
		offset_result.append(temp_value)
	print('offset_result',offset_result)
	return offset_result


#对偏离程度进行定性判断：高，偏高，正常范围，偏低，低，极端异常
def qualitative_offset(offset_result):
	#偏离程度定性标准，例如-20%~20%为正常，20%~35%为偏高，35%以上为高
	qualitative_standard=[0.2,0.35]
	qualitative_offset_result=[]
	for i in range(len(offset_result)):
		if abs(float(offset_result[i]))<=qualitative_standard[0]:
			qualitative_offset_result.append('在正常范围')
		elif abs(float(offset_result[i]))<=qualitative_standard[1]:
			if float(offset_result[i])>0:#偏高
				qualitative_offset_result.append('偏高')
			else:#偏低
				qualitative_offset_result.append('偏低')		
		else:
			if float(offset_result[i])>0:#高
				qualitative_offset_result.append('高')
			else:
				qualitative_offset_result.append('低')		
	return  qualitative_offset_result

#计算单炉次字段值
def single_heat(heat_no,fieldname):
	print("success")
	print(heat_no,fieldname)
	sqlVO1={}
	sqlVO1["db_name"]="l2own"
	sqlVO1["sql"]="SELECT HEAT_NO,nvl("+fieldname+",0) as field FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+heat_no+"'";
	print(sqlVO1["sql"])
	scrapy_records1=models.BaseManage().direct_select_query_sqlVO(sqlVO1)
	value = scrapy_records1[0].get('field',None)
	if value != None :
		scrapy_records1[0]['field'] = float(value)
	frame1=DataFrame(scrapy_records1)
	yaxis=frame1.FIELD[0]
	return yaxis

#根据时间段查询数据
def time(request):
	print("success")
	time1=request.POST.get("time1");
	time2=request.POST.get("time2");
	fieldname=request.POST.get("field").upper();
	print(time1)
	print(time2)
	sqlVO1={}
	sqlVO1["db_name"]="l2own"
	sqlVO1["sql"]="SELECT HEAT_NO,"+fieldname+",MSG_DATE_PLAN FROM qg_user.PRO_BOF_HIS_ALLFIELDS WHERE to_char(MSG_DATE_PLAN,'yyyy-mm-dd')>'"+time1+"' and to_char(MSG_DATE_PLAN,'yyyy-mm-dd')<'"+time2+"'";
	#sqlVO1["sql"]="SELECT HEAT_NO,"+fieldname+",MSG_DATE_PLAN FROM qg_user.PRO_BOF_HIS_ALLFIELDS WHERE to_char(MSG_DATE_PLAN,'yyyy-mm-dd')>'"+time1+"'";
	print(sqlVO1["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO1)
	print(scrapy_records[:5])
	contentVO={
		'title':'测试',
		'state':'success'
	}
	ana_result,ana_describe=num_describe(scrapy_records,fieldname)
	contentVO['result']=ana_result
	contentVO['describe']=ana_describe
	return HttpResponse(json.dumps(contentVO),content_type='application/json')


def quality_zhuanlu(request):
	#print('请求主页')
	if not request.user.is_authenticated():
		return HttpResponseRedirect("/login")
	return render(request,'data_import/qualityzhuanlu.html',{'title':"青特钢大数据项目组数据管理"})

#跳转波动率fluctuation.html页面
def fluctuation(request):
	if not request.user.is_authenticated():
		return HttpResponseRedirect("/login")
	return render(request,'data_import/fluctuation.html')

#从数据库动态加载钢种
def getGrape(request):
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="select distinct gk_no from qg_user.PRO_BOF_HIS_ALLFIELDS order by gk_no";
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	frame=DataFrame(scrapy_records)
	#print(frame['GK_NO'])
	contentVO={
		'title':'测试',
		'state':'success'
	}
	grape=[ele for ele in frame['GK_NO']]
	#print(grape)
	contentVO['result']=grape
	return HttpResponse(json.dumps(contentVO),content_type='application/json')

#进行五数分析法
def Wushu(x):
    L=np.percentile(x,25)-1.5*(np.percentile(x,75)-np.percentile(x,25))
    U=np.percentile(x,75)+1.5*(np.percentile(x,75)-np.percentile(x,25))
    result=x[(x<U)&(x>L)]
    wushu_clean={}
    wushu_clean["minbook"]=L
    wushu_clean["maxbook"]=U
    wushu_clean["result"]=result
    return wushu_clean

#数据清洗+计算概率分布（画概率直方图）
def num_describe(scrapy_records,bookno):
	print(scrapy_records[1:5])
	ivalue_i=[]
	for n in range(len(scrapy_records)):
		ivalue = scrapy_records[n].get(bookno,None)
		if ivalue !=None and ivalue !=0:
			 ivalue=ivalue
			 ivalue_b=str(ivalue)
			 ivalue_valid=ivalue_num(ivalue_b)#小数位数
			 ivalue_i.append(ivalue_valid)
			 ivalue_i.sort(reverse=True)
	ivalue_valid=ivalue_i[0]#取所有有效位数的最大个数
	print(ivalue_valid)		
	for i in range(len(scrapy_records)):
		value = scrapy_records[i].get(bookno,None)
		if value != None :
			scrapy_records[i][bookno] = float(value)			
	frame=DataFrame(scrapy_records)	
	df=frame.sort_values(by=bookno)
	dfr=df[df>0].dropna(how='any')
	#print(dfr['1622324'])
	#print(dfr[bookno].dtype)
	cleanbook=Wushu(dfr[bookno])
	minbook=cleanbook["minbook"]
	maxbook=cleanbook["maxbook"]	
	if(minbook==maxbook):#不符合正态分布规律
		print("no")
		clean=dfr[bookno]		
	else:
		print("yes")
		clean=cleanbook["result"]				
	print("minbook")
	print(minbook)
	print("maxbook")
	print(maxbook)
	print(type(clean))
	if clean is not None:
		if(clean.max==clean.min()):
			bc=1
		else:	
			bc=(clean.max()-clean.min())/7
		bcq=math.ceil(bc*1000)/1000
		print(bcq)
		try:
			section=pd.cut(clean,math.ceil((clean.max()-clean.min())/bcq))
			end=pd.value_counts(section,sort=False)/clean.count()
			describe=clean.describe()
		except ValueError as e:
		 	print(e)
	numx=[ele for ele in end.index]
	#numy=[ele for ele in end]
	desx=[ele for ele in describe.index]
	desy=[ele for ele in describe]
	print(numx)
	print(describe)
	#contentVO={
		#'title':'测试',
		#'state':'success'
	#}
	s1=pd.Series(numx)
	d1=getA(s1)
	d2=getB(s1)
	d1_data=[]
	d2_data=[]
	#d3_data=[]
	d4_data=[]
	d1_valid=vaild(d1,ivalue_valid,d1_data)
	for i in range(len(d1_valid)):
		if d1_valid[i]<0:
			d1_valid[i]=0
	d2_valid=vaild(d2,ivalue_valid,d2_data)
	numx1=list(set(d1_valid).union(set(d2_valid)))
	numx2=sorted(numx1)
	print("numx2:")
	print(numx2)
	sections=[]
	numx3=union_section(numx2,sections)
	cut1=pd.cut(clean,numx2)
	end1=pd.value_counts(cut1,sort=False)/clean.count()
	numy=[ele for ele in end1]
	print("end1:")
	print(end1)
	#numy1=vaild(numy,ivalue_valid,d3_data)
	numy1=["%.6f"%(n) for n in numy]
	desy1=vaild(desy,ivalue_valid,d4_data)
	ana_result={}
	ana_result['scope']=numx3
	#print(ana_result['scope'])
	ana_result['num']=numy1
	#contentVO['result']=ana_result
	ana_describe={}
	ana_describe['scopeb']=desx
	ana_describe['numb']=desy1
	#contentVO['describe']=ana_describe
	return ana_result,ana_describe

#判断有效位数
def ivalue_num(num):
    a=str(num)
    if(a.isdigit()):
        ivalue_valid=0
    else:
        ivalue_valid=len(a.split('.')[1])
    return  ivalue_valid   

#取有效位数
def vaild(lis,ivalue_valid,data):
    for i in range(len(lis)):
        shu=lis[i]
        if ivalue_valid==0:
            shua=int(float(shu))
            data.append(shua)
        else:    
            shua=float(shu)
            shub=round(shua,ivalue_valid)
            data.append(shub)
    return data 

#数据清洗+计算正态分布（画正态分布图）
#与num_describe的区别是，进行数据清洗并计算正态分布，而不进行概率直方图区间计算
from scipy.stats import norm
def data_clean(scrapy_records,bookno):
	print("data_clean:"+bookno);
	if bookno=='"AS"':
		bookno=bookno.split('"')[1]
	for i in range(len(scrapy_records)):
		value = scrapy_records[i].get(bookno,None)
		if value != None :
			scrapy_records[i][bookno] = float(value)
	frame=DataFrame(scrapy_records)	#可能有多列数据
	frame_formal=frame[['HEAT_NO',bookno]]
	#print(frame[bookno])
	df=frame.sort_values(by=bookno)
	dfr=df[df>0].dropna(how='any')
	#print(dfr[bookno].dtype)
	clean=Wushu(dfr[bookno])
	if bookno=="NB":
		print(clean)
	#print("clean",clean)
	#平均值/期望值
	dataclean_result={}
	avg_value=np.mean(clean)
	print("期望值",avg_value)
	#标准差
	std_value=np.std(clean)
	print("标准差",std_value)
	#方差
	var_value=np.var(clean)
	print("方差",var_value)
	#normx,normy=Norm_dist(avg_value,var_value)
	print("min",clean.min())
	print("max",clean.max())
	aa=(clean.max()-clean.min())/50
	normx = np.arange(clean.min(),clean.max(),aa)  
	normy = norm.pdf(normx,avg_value,std_value)
	dataclean_result['clean_min']=clean.min()
	dataclean_result['clean_max']=clean.max()
	dataclean_result['avg_value']=avg_value#期望值
	dataclean_result['std_value']=std_value#标准差
	dataclean_result['normx']=normx#x轴取样点
	dataclean_result['normy']=normy#正态分布取样点对应的Y轴取值
	return dataclean_result

#同时计算概率分布和正态分布
from scipy.stats import norm
def probability_distribution(request):
	print("同时计算概率直方图和正态分布图！")

	#获取信息
	heat_no=request.POST.get("heat_no");#炉次号
	bookno=request.POST.get("fieldname_english").upper();#字段英文名
	fieldname_chinese=request.POST.get("fieldname_chinese")#字段中文名
	offset_value=request.POST.get("offset_value");#偏离值
	qualitative_offset_result=request.POST.get("qualitative_offset_result");#定性判断
	# offset_value=float(offset_value_temp[1:-1])/100
	actual_value=float(request.POST.get("actual_value"))#实际值
	coloum_number=int(request.POST.get("coloum_number"))
	print(heat_no,bookno,actual_value,coloum_number);
	# print(scrapy_records[1:5])
	str_select=str(request.POST.get("str_select"));#筛选条件
	if str_select==str(None):#表示没有筛选条件
		str_select=''

	#查询上下限及是否五数清洗
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT NUMERICAL_LOWER_BOUND,NUMERICAL_UPPER_BOUND,IF_FIVENUMBERSUMMARY FROM qg_user.PRO_BOF_HIS_ALLSTRUCTURE where DATA_ITEM_EN = \'"+bookno+"\'"
	# print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	NUMERICAL_LOWER_BOUND=scrapy_records[0].get('NUMERICAL_LOWER_BOUND',None)
	NUMERICAL_UPPER_BOUND=scrapy_records[0].get('NUMERICAL_UPPER_BOUND',None)
	IF_FIVENUMBERSUMMARY=scrapy_records[0].get('IF_FIVENUMBERSUMMARY',None)

	bound_select =' and '+bookno+'>='+str(scrapy_records[0]['NUMERICAL_LOWER_BOUND']) +' and '+bookno+'<='+str(scrapy_records[0]['NUMERICAL_UPPER_BOUND'])

	#需要返回的计算结果
	contentVO={
	'title':'测试',
	'state':'success'
	}

	#进行数据库查询
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,"+bookno+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS WHERE HEAT_NO>'1500000'"+str_select+bound_select
	#sqlVO["sql"]="SELECT HEAT_NO,"+bookno+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS"
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	if len(scrapy_records)<10:
		contentVO['state']='error'
		return HttpResponse(json.dumps(contentVO),content_type='application/json')

	#进行数据清洗
	ivalue_i=[]
	for n in range(len(scrapy_records)):
		ivalue = scrapy_records[n].get(bookno,None)
		if ivalue !=None and ivalue !=0:
			 ivalue=ivalue
			 ivalue_b=str(ivalue)
			 ivalue_valid=ivalue_num(ivalue_b)#小数位数
			 ivalue_i.append(ivalue_valid)
			 ivalue_i.sort(reverse=True)
	ivalue_valid=ivalue_i[0]#取所有有效位数的最大个数
	print(ivalue_valid)		
	for i in range(len(scrapy_records)):
		value = scrapy_records[i].get(bookno,None)
		if value != None :
			scrapy_records[i][bookno] = float(value)			
	frame=DataFrame(scrapy_records)	
	df=frame.sort_values(by=bookno)
	dfr=df[df>0].dropna(how='any')
	#print(dfr['1622324'])
	#print(dfr[bookno].dtype)
	cleanbook=Wushu(dfr[bookno])
	minbook=cleanbook["minbook"]
	maxbook=cleanbook["maxbook"]	
	if(minbook==maxbook):#不符合正态分布规律
		print("no")
		clean=dfr[bookno]		
	else:
		print("yes")
		clean=cleanbook["result"]				
	print("minbook")
	print(minbook)
	print("maxbook")
	print(maxbook)
	print(type(clean))

	#计算概率直方分布起始——————————————————————————————————————————————————————————————————————————————————————————————
	if clean is not None:
		if(clean.max==clean.min()):
			bc=1
		else:	
			bc=(clean.max()-clean.min())/coloum_number
		bcq=math.ceil(bc*1000)/1000
		print(bcq)
		try:
			section=pd.cut(clean,math.ceil((clean.max()-clean.min())/bcq))
			end=pd.value_counts(section,sort=False)/clean.count()
			describe=clean.describe()
		except ValueError as e:
		 	print(e)
	numx=[ele for ele in end.index]
	#numy=[ele for ele in end]
	desx=[ele for ele in describe.index]
	desy=[ele for ele in describe]
	print(numx)
	print(describe)
	#contentVO={
		#'title':'测试',
		#'state':'success'
	#}
	s1=pd.Series(numx)
	d1=getA(s1)
	d2=getB(s1)
	d1_data=[]
	d2_data=[]
	#d3_data=[]
	d4_data=[]
	d1_valid=vaild(d1,ivalue_valid,d1_data)
	for i in range(len(d1_valid)):
		if d1_valid[i]<0:
			d1_valid[i]=0
	d2_valid=vaild(d2,ivalue_valid,d2_data)
	numx1=list(set(d1_valid).union(set(d2_valid)))
	numx2=sorted(numx1)
	# print("numx2:")
	# print(numx2)
	sections=[]
	numx3=union_section(numx2,sections)
	cut1=pd.cut(clean,numx2)
	end1=pd.value_counts(cut1,sort=False)/clean.count()
	numy=[ele for ele in end1]
	# print("end1:")
	# print(end1)
	#numy1=vaild(numy,ivalue_valid,d3_data)
	numy1=["%.6f"%(n) for n in numy]
	print("概率分布：直方图bar")
	print(numx3)
	print(numy1)
	desy1=vaild(desy,ivalue_valid,d4_data)
	ana_result={}
	ana_result['scope']=numx3#区间数组
	#print(ana_result['scope'])
	ana_result['num']=numy1#区间对应值
	#contentVO['result']=ana_result
	ana_describe={}
	ana_describe['scopeb']=desx
	ana_describe['numb']=desy1
	#contentVO['describe']=ana_describe
	# return ana_result,ana_describe
	#计算概率直方分布结束——————————————————————————————————————————————————————————————————————————————————————————————
	#计算正态分布起始——————————————————————————————————————————————————————————————————————————————————————————————————
	
	avg_value=np.mean(clean)
	print("期望值",avg_value)
	#标准差
	std_value=np.std(clean)
	print("标准差",std_value)
	#方差
	var_value=np.var(clean)
	print("方差",var_value)
	#normx,normy=Norm_dist(avg_value,var_value)
	print("min",clean.min())
	print("max",clean.max())
	aa=(clean.max()-clean.min())/50
	normx_array = np.arange(clean.min(),clean.max(),aa)  
	normy_array = norm.pdf(normx_array,avg_value,std_value)
	#转换格式，限制小数位数
	normx=["%.4f"%(n) for n in list(normx_array)]
	normy=["%.7f"%(n) for n in list(normy_array)]
	print("正态分布：曲线line")
	print("normx",normx)
	print("normy",normy)
	normal_result={}
	normal_result['clean_min']=clean.min()#最小值
	normal_result['clean_max']=clean.max()#最大值
	normal_result['avg_value']=avg_value#期望值
	normal_result['std_value']=std_value#标准差
	normal_result['normx']=normx#x轴取样点
	normal_result['normy']=normy#正态分布取样点对应的Y轴取值
	normal_result['fieldname']=bookno#字段英文名
	normal_result['fieldname_chinese']=fieldname_chinese#字段中文名
	normal_result['actual_value']=actual_value#自身值
	# normal_result['offset_value']="%.2f%%"%(offset_value*100)#偏离程度(当offset_value为小数形式时可用此行代码)
	normal_result['offset_value']=offset_value#偏离程度(当offset_value已经为百分号形式时可用此行代码)
	# return normal_result

	#---------判断该炉次该字段在历史正态分布曲线中距离最近的取样点-----------
	temp=float(normx[0])
	temp_index=0
	for i in range (0,len(normx)-1):
		former=float(normx[i])
		latter=float(normx[i+1])
		#print(former,latter)
		if (actual_value>former) & (actual_value<=latter):
			temp_index=i
			if abs(actual_value-former)<abs(actual_value-latter):
				temp=former
			else:
				temp=latter
			break;
	print(temp_index,temp)

	normal_result['match_index']=temp_index#对应序号
	normal_result['match_value']=str("%.4f"%(temp))#距离最近的值,由于在之前将其转化为float类型进行过数据对比，例如，123.0000被默认识别为123，当再次转化为str类型时，则也成了‘123’，因此就与normx的值对不上了
	
	#--------------------------------------------------------------------------
	#计算正态分布结束———————————————————————————————————————————————————————————————————————————————————————————————————

	#返回概率分布及正态分布计算结果
	final_result={}
	final_result['ana_result']=ana_result#概率分布
	final_result['ana_describe']=ana_describe
	final_result['normal_result']=normal_result#正态分布
	final_result['qualitative_offset_result']=qualitative_offset_result#定性判断
	# return final_result
	return HttpResponse(json.dumps(final_result),content_type='application/json')
	
#影响因素追溯
from . import zhuanlu
def max_influence(request):
	print("哈爽")
	print("Enter max_influence")
	field=request.POST.get("field");
	offset_value=request.POST.get("offset_value");#炉次字段的偏离程度，
	print(field)
	#将待匹配的字段名设置为参数
	#field='MIRON_WGT'
	result=[]
	#从数据库读取回归系数文件
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT * FROM pro_bof_his_REGRESSION_COF where MIDFIELD='"+field +"' and INFIELD != 'BIAS' order by abs(COF) desc"
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	print(scrapy_records)
	length_result1=len(scrapy_records)

	xasis_fieldname=[]#字段英文名字数组
	regression_coefficient=[]#字段回归系数值数组
	str_sql=''
	for i in range(length_result1):
		xasis_fieldname.append(scrapy_records[i].get('INFIELD', None))
		regression_coefficient.append(scrapy_records[i].get('COF', None))
		str_sql=str_sql+','+scrapy_records[i].get('INFIELD', None)
	prime_cost=request.POST.get("prime_analyse");
	sqlVO={}
	sqlVO["db_name"]="l2own"
	# sqlVO["sql"]="SELECT HEAT_NO,nvl("+xasis_fieldname[0]+",0) as "+xasis_fieldname[0]+",nvl("+xasis_fieldname[1]+",0) as "+xasis_fieldname[1]+",nvl("+xasis_fieldname[2]+",0) as "+xasis_fieldname[2]+",nvl("+xasis_fieldname[3]+",0) as "+xasis_fieldname[3]+",nvl("+xasis_fieldname[4]+",0) as "+xasis_fieldname[4]+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+prime_cost+"'";
	sqlVO["sql"]="SELECT HEAT_NO" +str_sql+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where HEAT_NO='"+prime_cost+"'";
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	#将查询所得值全部转变为float格式（动态，适用于不同个数的xasis_fieldname长度）
	yaxis=[]#各相关性字段的实际值
	for i in range(length_result1):
		value = scrapy_records[0].get(xasis_fieldname[i],None)
		if value != None :
			scrapy_records[0][xasis_fieldname[i]] = Decimal(value)
		else:
			scrapy_records[0][xasis_fieldname[i]] = 0#将空值None暂时以0填充
		yaxis.append(value)
	frame=DataFrame(scrapy_records)
	print("frame",frame)
	# for i in range(length_result1):
	# 	yaxis.append(frame[xasis_fieldname[i]][0])
	# yaxis=[frame[xasis_fieldname[0]][0],frame[xasis_fieldname[1]][0],frame[xasis_fieldname[2]][0],frame[xasis_fieldname[3]][0],frame[xasis_fieldname[4]][0]]
	print("yaxis",yaxis)
	contentVO={
		'title':'测试',
		'state':'success'
	}
	offset_result=offset(xasis_fieldname,yaxis)
	print("xasis_fieldname",xasis_fieldname)
	print("偏离程度")
	print(offset_result)
	xasis_fieldname_result=[]#字段英文名字数组
	regression_coefficient_result=[]#字段回归系数值数组
	offset_result_final=[]#偏离程度值

	#即字段偏离高，则正相关应对应偏高，负相关应对应偏低
	if float(offset_value[0:-1])>=0:#读取隐藏域的值，由于隐藏域的偏离表示为百分比，例如12.6%。因此截取12.6来判断其正负
		for i in range(length_result1):
			if offset_result[i]==None or xasis_fieldname[i]=='NB':#由于数据清洗的问题，暂且将NB字段如此处理，因为NB字段的所有数据均相同，导致数据清洗时将所有数据都清除了
				continue
			if float(regression_coefficient[i]) * float(offset_result[i]) >=0:
				xasis_fieldname_result.append(xasis_fieldname[i])
				regression_coefficient_result.append(regression_coefficient[i])
				offset_result_final.append(offset_result[i])
	else:
		for i in range(length_result1):
			if offset_result[i]==None or xasis_fieldname[i]=='NB':
				continue
			if float(regression_coefficient[i]) * float(offset_result[i]) <=0:
				xasis_fieldname_result.append(xasis_fieldname[i])
				regression_coefficient_result.append(regression_coefficient[i])
				offset_result_final.append(offset_result[i])
	contentVO['xasis_fieldname']=xasis_fieldname_result#回归系数因素英文字段名
	contentVO['regression_coefficient']=regression_coefficient_result#字段回归系数值
	contentVO['offset_result']=offset_result_final#回归系数因素字段偏离值
	contentVO['offset_number']=len(xasis_fieldname_result)#回归系数最大因素所取字段个数
	print("字段名字")
	print(xasis_fieldname_result)
	print("回归系数")
	print(regression_coefficient_result)
	print("偏离程度")
	print(offset_result_final)
	#相关字段进行简单定性判断
	offset_result_nature=simple_offset(offset_result_final)
	print("简单定性判断")
	print(offset_result_nature)
	contentVO['offset_result_nature']=offset_result_nature

	ana_result={}
	ana_result=zhuanlu.PRO_BOF_HIS_ALLFIELDS
	En_to_Ch_result=[]
	for i in range(len(xasis_fieldname_result)):
		En_to_Ch_result.append(ana_result[xasis_fieldname_result[i]])
	print("中文字段")	
	print(En_to_Ch_result)

	ana_result_score={}
	ana_result_score=zhuanlu.PRO_BOF_HIS_ALLFIELDS_SCORE
	En_to_Ch_result_score=[]
	for i in range(len(xasis_fieldname_result)):
		En_to_Ch_result_score.append(ana_result_score[xasis_fieldname_result[i]])
	print("带标记的中文字段")	
	print(En_to_Ch_result_score)	

	#取前5个权重最大的字段按操作顺序（表结构）进行排序
	#中文字段与偏离程度合并，需联动排序
	offset_result_final_maxfive=offset_result_final[0:5]
	En_to_Ch_result_maxfive=En_to_Ch_result[0:5]
	dicty=dict(zip(En_to_Ch_result_maxfive,offset_result_final_maxfive))
	print('组合字段中文名和偏离程度')
	print(dicty)

	#按操作排序
	a=[]
	En_to_Ch_result_score=[]
	xasis_fieldname_result_maxfive=xasis_fieldname_result[0:5]
	for i in range(len(xasis_fieldname_result_maxfive)):
		a.append(ana_result_score[xasis_fieldname_result_maxfive[i]])
	# print('带标记的字段')
	# print(a)
	L=sorted(a,key=by_score)
	for i in range(len(xasis_fieldname_result_maxfive)):
		En_to_Ch_result_score.append(L[i][0])
	print('操作排序后中文名')
	print(En_to_Ch_result_score)

    #联动排序偏离程度
	offset_result_final_maxfive_order=[]
	for i in range(len(En_to_Ch_result_score)):
		offset_result_final_maxfive_order.append(dicty[En_to_Ch_result_score[i]])
	print('按操作排序后字段偏离程度')
	print(offset_result_final_maxfive_order)	
    #联动排序偏离程度
	# dictx=sorted(dicty.items(),key=lambda d:d in b )	
	# print('联动排序')
	# print(dictx)
	# offset_result_final_maxfive_order=[dictx[i][1] for i in range(len(dictx))]#按操作排序后字段中文名
	# xasis_fieldname_result_maxfive_order=[dictx[i][0] for i in range(len(dictx))]#按操作排序后字段偏离程度
	# print('按操作排序后字段英文名')
	# print(xasis_fieldname_result_maxfive_order)
	
	
	pos=map(lambda x:abs(x),offset_result_final_maxfive_order)
	posNum=["%.2f%%"%(n*100) for n in list(pos)]
	print("格式化")
	print(posNum)
	contentVO['En_to_Ch_result']=En_to_Ch_result#回归系数最大因素中文字段名字

	#contentVO['offset_result_final_maxfive_order']=offset_result_final_maxfive_order#按操作排序后字段偏离程度
	contentVO['En_to_Ch_result_score']=En_to_Ch_result_score#按操作排序后字段中文名
	contentVO['posNum']=posNum#按操作排序后偏离程度格式话

	return HttpResponse(json.dumps(contentVO),content_type='application/json')
def by_score(t):
    return t[1]	

#更新数据库转炉表结构期望等参数：updatevalue+Calculation_Parameters-----------------------------------------------------------------------------------------------------------------
import cx_Oracle
#定期更新数据库转炉字段统计值
def updatevalue(request):
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="select DATA_ITEM_EN,IF_ANALYSE,IF_FIVENUMBERSUMMARY from PRO_BOF_HIS_ALLSTRUCTURE"#读取字段列表,三个字段分别为字段名，是否用于分析，是否进行五值计算
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	print(len(scrapy_records))
	tns=cx_Oracle.makedsn('202.204.54.75',1521,'orcl')
	db=cx_Oracle.connect('qg_user','123456',tns)
	cur = db.cursor()#创建cursor
	# sql_str="select DATA_ITEM_EN,IF_ANALYSE,IF_FIVENUMBERSUMMARY from PRO_BOF_HIS_ALLSTRUCTURE"#读取字段列表,三个字段分别为字段名，是否用于分析，是否进行五值计算
	# cur.execute(sql_str)
	# rs=cur.fetchall()  #一次返回所有结果集
	for rs in scrapy_records:
		print(rs);
		# print(rs["DATA_ITEM_EN"],rs["IF_ANALYSE"],rs["IF_FIVENUMBERSUMMARY"])
		#符合分析条件则进行极值及期望标准差计算
		if rs["IF_ANALYSE"] != '1':
			continue;
		#计算参数
		dataclean_result=Calculation_Parameters(rs["DATA_ITEM_EN"],rs["IF_FIVENUMBERSUMMARY"]);
		min_value=str(dataclean_result["clean_min"])#最小值
		max_value=str(dataclean_result["clean_max"])#最大值
		average_value=str(dataclean_result['avg_value'])#期望值
		standard_value=str(dataclean_result['std_value'])#标准差

		#更新数据库
		sql_str="UPDATE PRO_BOF_HIS_ALLSTRUCTURE SET MAX_VALUE ="+max_value+", MIN_VALUE="+min_value+", DESIRED_VALUE="+average_value+", STANDARD_DEVIATION="+standard_value+" WHERE DATA_ITEM_EN = \'"+rs["DATA_ITEM_EN"]+"\'";
		print(sql_str)
		try:
			cur.execute(sql_str)
		except:
			print(rs["DATA_ITEM_EN"]+"update failed!")
			pass
		db.commit()

	# for item in data:
	#     for i in range(len(item)): 
	#         str_item=str(item[i])
	#         sql_str_temp=str_item[1:-1]
	#         sql_str="insert into PRO_BOF_HIS_ALLSTRUCTURE values("+sql_str_temp+")"
	#         print(sql_str)
	#         try:
	#              c.execute(sql_str)
	#             # db.commit()
	#         except:
	#             print('insert PRO_BOF_HIS_ALLSTRUCTURE failed: ',i)
	#             # print(sql_str)
	#             pass
	#         # print(i)
	#         # print(item[i])
	cur.close()
	# db.commit()
	db.close()
	contentVO={
		'title':'测试',
		'state':'success'
	}
	return HttpResponse(json.dumps(contentVO),content_type='application/json')

#计算极值、期望和标准差
from scipy.stats import norm
def Calculation_Parameters(fieldname,IF_FIVENUMBERSUMMARY):
	# print("enter Calculation_Parameters!")
	#对字段AS的特殊情况进行处理
	if fieldname=='AS':
		fieldname='"AS"'
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,"+fieldname+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS"
	print(sqlVO["sql"])
	scrapy_records=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	#print(fieldname)
	#print(scrapy_records[:5])
	contentVO={
		'title':'测试',
		'state':'success',
	}
	print("data_clean:"+fieldname);
	if fieldname=='"AS"':
		fieldname=fieldname.split('"')[1]

	for i in range(len(scrapy_records)):
		value = scrapy_records[i].get(fieldname,None)
		if value != None :
			scrapy_records[i][fieldname] = float(value)
	frame=DataFrame(scrapy_records) #可能有多列数据
	# frame_formal=frame[['HEAT_NO',fieldname]]
	#print(frame[fieldname])
	#字段NB的所有数据都相同，因此列为特殊情况，不对其进行清洗处理
	df=frame.sort_values(by=fieldname)
	dfr=df[df>0].dropna(how='any')
	#print(dfr[fieldname].dtype)
	if IF_FIVENUMBERSUMMARY=='1':#是否进行五数分析
		print("进行五值分析:",fieldname)
		clean=Wushu(dfr[fieldname])["result"]
	else:
		clean=dfr[fieldname]

	#print("clean",clean)
	#平均值/期望值
	dataclean_result={}
	avg_value=np.mean(clean)
	print("期望值",avg_value)
	#标准差
	std_value=np.std(clean)
	print("标准差",std_value)
	#方差
	var_value=np.var(clean)
	print("方差",var_value)
	#normx,normy=Norm_dist(avg_value,var_value)
	print("min",clean.min())
	print("max",clean.max())

	dataclean_result['clean_min']=clean.min()
	dataclean_result['clean_max']=clean.max()
	dataclean_result['avg_value']=avg_value#期望值
	dataclean_result['std_value']=std_value#标准差

	return dataclean_result
#对相关性字段回归系数展示做简单定性分析	正数为偏高负数为偏低
def simple_offset(offset_result_final):
	offset_result_nature=[]#相关字段定性分析
	for i in range(len(offset_result_final)):
		if float(offset_result_final[i]>0):
			offset_result_nature.append('偏高')
		else:
			offset_result_nature.append('偏低')
	return 	offset_result_nature		


#暴力求解
from . import batchprocess
'''
单炉次字段质量回溯
网页
'''
def retrospectfactor_all(request):
	prime_cost=request.POST.get("prime_cost");
	str_select=str(request.POST.get("str_select"));	
	print(prime_cost)
	str_all=retrospectfactor_all_to(prime_cost,str_select)
		
	contentVO={
		'title':'测试',
		'state':'success'		
	}
	contentVO['str_all']=str_all				
	return HttpResponse(json.dumps(contentVO),content_type='application/json')
def retrospectfactor_all_to(prime_cost,str_select):
	#取某一炉次的质量分析字段C、SI、MN、P、S、钢水重量、温度的实际值
	
	str_all=" "
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,nvl(C,0) as C,nvl(SI,0) as SI ,nvl(MN,0) as MN,nvl(P,0) as P ,nvl(S,0) as S, nvl(FE,0) as FE, nvl(STEELWGT,0) as STEELWGT,nvl(FINAL_TEMP_VALUE,0)as FINAL_TEMP_VALUE FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+prime_cost+"'";
	
	# sqlVO["sql"] = "SELECT " + ",".join(xasis_fieldname_single) + " from QG_USER.PRO_BOF_HIS_ALLFIELDS where heat_no= '"+prime_cost+"'"

	scrapy_records_single=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	xasis_fieldname_single=['C','SI','MN','P','S','STEELWGT','FINAL_TEMP_VALUE']
	xaxis=['C含量','Si含量','Mn含量','P含量','S含量','重量','温度']

	
	str_heatno='炉次号'+prime_cost+''      
	str_all=str_all+str_heatno+'\n'
	
	for i in range(len(xasis_fieldname_single)):
		value = scrapy_records_single[0].get(xasis_fieldname_single[i],None)
		if value != None :
			scrapy_records_single[0][xasis_fieldname_single[i]] = float(value)
			#yaxis_single.append(float(value))
	frame=DataFrame(scrapy_records_single)
	#实际值
	yaxis_single=[frame.C[0],frame.SI[0],frame.MN[0],frame.P[0],frame.S[0],frame.STEELWGT[0],frame.FINAL_TEMP_VALUE[0]]
	#yaxis_single=[frame.C[0]]

	print('单炉次质量分析字段')
	print(xasis_fieldname_single)
	print("单炉次质量分析字段实际值")
	print(yaxis_single)
	
	#计算单炉次质量分析字段偏离程度
	offset_result_single=[]
	offset_result_single_a=offset(xasis_fieldname_single,yaxis_single,str_select)
	#介于五数概括法后的极值与规定的范围之间的数即偏离程度超过50%的按50%处理	
	for n in range(len(offset_result_single_a)):
		if abs(float(offset_result_single_a[n]))>0.5:
			if float(offset_result_single_a[n])<0:
				offset_result_single.append('-0.5')
			else:	
				offset_result_single.append('0.5')
		else:
			offset_result_single.append(offset_result_single_a[n])

	offset_value_single=["%.2f%%"%(n*100) for n in list(offset_result_single)]
	print('单炉次质量分析字段偏离度')
	print(offset_value_single)

	#分析字段偏离程度定性判断
	qualitative_offset_result=qualitative_offset(offset_result_single)
	print('单炉次质量分析字段偏离度定向分析')
	print(qualitative_offset_result)
	#写入word

	
	for i in range(len(xasis_fieldname_single)):
		xaxis_chinese=xaxis[i];
		field=xasis_fieldname_single[i];
		single_value=yaxis_single[i];
		offset_value=offset_value_single[i];
		qualitative_offset_result_single=qualitative_offset_result[i];
		En_to_Ch_result_score,offset_result_nature,offset_value_single_cof,coef_value_re=analy_cof(prime_cost,field,single_value,offset_value);				
		
		#偏离程度小于20%
		if abs(float(offset_result_single[i]))<=0.2:
			continue;
		#超过上下限范围	
		elif bound_judge(field,single_value)==0:
			continue;
		#回归分析后无数据	
		elif En_to_Ch_result_score==None:
		 	continue;		
		else:
			str_des=xaxis_chinese+qualitative_offset_result_single+',偏离度为'+offset_value+'原因是:'      
			str_all=str_all+str_des+'\n'	
			for i in range(len(En_to_Ch_result_score)):
				if i<len(En_to_Ch_result_score)-1:
					if float(offset_value_single_cof[i][0:-1])>50:
						str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+'50%'+','
					else:
						str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+','
				else:
					if float(offset_value_single_cof[i][0:-1])>50:
						str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+'50%'+'\n'
					else:
						str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+'\n'	
				
				str_all=str_all+str_cause+''
    
	return 	str_all			
'''
单炉次字段质量回溯
暴力求解
'''
def violent_ananlyse(request):
	print("violent_ananlyse")
	document = Document()
	paragraph=document.add_paragraph()
	#炉次	
	for i in range(1):
		# prime_cost=str(1530320+i);
		prime_cost=str(1634230+i);		
		str_cause=violent_ananlyse_to(prime_cost,document,paragraph)
	
	contentVO={
		'title':'测试',
		'state':'success'		
	}
	contentVO['result']=ana_result				
	return HttpResponse(json.dumps(contentVO),content_type='application/json')	

def violent_ananlyse_to(prime_cost,document,paragraph):
	#取某一炉次的质量分析字段C、SI、MN、P、S、钢水重量、温度的实际值
	

	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO,nvl(C,0) as C,nvl(SI,0) as SI ,nvl(MN,0) as MN,nvl(P,0) as P ,nvl(S,0) as S, nvl(FE,0) as FE, nvl(STEELWGT,0) as STEELWGT,nvl(FINAL_TEMP_VALUE,0)as FINAL_TEMP_VALUE FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+prime_cost+"'";
	
	# sqlVO["sql"] = "SELECT " + ",".join(xasis_fieldname_single) + " from QG_USER.PRO_BOF_HIS_ALLFIELDS where heat_no= '"+prime_cost+"'"

	scrapy_records_single=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	xasis_fieldname_single=['C','SI','MN','P','S','FE','STEELWGT','FINAL_TEMP_VALUE']
	xaxis=['C含量','SI含量','MN含量','P含量','S含量','FE含量','重量','温度']

	
	str_heatno='炉次号'+prime_cost+'\n'      
	paragraph.add_run(str_heatno)

	# xasis_fieldname_single=['C']
	# xaxis=['C含量']
	
	for i in range(len(xasis_fieldname_single)):
		value = scrapy_records_single[0].get(xasis_fieldname_single[i],None)
		if value != None :
			scrapy_records_single[0][xasis_fieldname_single[i]] = float(value)
			#yaxis_single.append(float(value))
	frame=DataFrame(scrapy_records_single)
	#实际值
	yaxis_single=[frame.C[0],frame.SI[0],frame.MN[0],frame.P[0],frame.S[0],frame.FE[0],frame.STEELWGT[0],frame.FINAL_TEMP_VALUE[0]]
	#yaxis_single=[frame.C[0]]

	print('单炉次质量分析字段')
	print(xasis_fieldname_single)
	print("单炉次质量分析字段实际值")
	print(yaxis_single)
	
	#计算单炉次质量分析字段偏离程度
	offset_result_single=[]
	offset_result_single_a=offset(xasis_fieldname_single,yaxis_single)
	#介于五数概括法后的极值与规定的范围之间的数即偏离程度超过50%的按50%处理	
	for n in range(len(offset_result_single_a)):
		if abs(float(offset_result_single_a[n]))>0.5:
			if float(offset_result_single_a[n])<0:
				offset_result_single.append('-0.5')
			else:	
				offset_result_single.append('0.5')
		else:
			offset_result_single.append(offset_result_single_a[n])

	offset_value_single=["%.2f%%"%(n*100) for n in list(offset_result_single)]
	print('单炉次质量分析字段偏离度')
	print(offset_value_single)

	#分析字段偏离程度定性判断
	qualitative_offset_result=qualitative_offset(offset_result_single)
	print('单炉次质量分析字段偏离度定向分析')
	print(qualitative_offset_result)
	#写入word

	
	for i in range(len(xasis_fieldname_single)):
		xaxis_chinese=xaxis[i];
		field=xasis_fieldname_single[i];
		single_value=yaxis_single[i];
		offset_value=offset_value_single[i];
		qualitative_offset_result_single=qualitative_offset_result[i];
		En_to_Ch_result_score,offset_result_nature,offset_value_single_cof,coef_value_re=analy_cof(prime_cost,field,single_value,offset_value);				
		
		#偏离程度小于20%
		if abs(float(offset_result_single[i]))<=0.2:
			continue;
		#超过上下限范围	
		elif bound_judge(field,single_value)==0:
			continue;
		#回归分析后无数据	
		elif En_to_Ch_result_score==None:
		 	continue;	
		else:
			str_des=xaxis_chinese+qualitative_offset_result_single+',偏离度为'+offset_value+'原因是:'      
			paragraph.add_run(str_des)	
			for i in range(len(En_to_Ch_result_score)):
				if i<len(En_to_Ch_result_score)-1:
					str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+','
				else:
					str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+'\n'	
				paragraph.add_run(str_cause) 
    
	document.add_page_break()
	document.save('F:/demo.docx')
	return 	ana_result		
	
def  analy_cof(prime_cost,field,single_value,offset_value):
	'''
	prime_cost：炉次号
	field：单炉次质量分析字段
	single_value：单炉次质量分析字段实际值
	offset_value：单炉次质量分析字段偏离值
	'''
	#从数据库读取相关字段并按照相关系数绝对值由大到小排序
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT * FROM pro_bof_his_relation_cof where OUTPUTFIELD='"+field +"'order by abs(COF) desc"
	#print(sqlVO["sql"])
	scrapy_records_relation=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	#print(scrapy_records)
	length_result1=len(scrapy_records_relation)

	xasis_fieldname_start=[]#字段英文名字数组
	regression_relation_start=[]#字段相关系数值数组
	str_sql=''
	for i in range(length_result1):
		xasis_fieldname_start.append(scrapy_records_relation[i].get('MIDDLEFIELD', None))
		regression_relation_start.append(scrapy_records_relation[i].get('COF', None))
		str_sql=str_sql+','+scrapy_records_relation[i].get('MIDDLEFIELD', None)
	
	#取相关字段实际值
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO" +str_sql+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where HEAT_NO='"+prime_cost+"'";
	print(sqlVO["sql"])
	scrapy_records_actual=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	yaxis=[]#各相关性字段实际值
	for i in range(length_result1):
		value = scrapy_records_actual[0].get(xasis_fieldname_start[i],None)
		if value != None :
			scrapy_records_actual[0][xasis_fieldname_start[i]] = Decimal(value)
		else:
			scrapy_records_actual[0][xasis_fieldname_start[i]] = 0#将空值None暂时以0填充
		yaxis.append(value)
	frame=DataFrame(scrapy_records_actual)
	print("frame",frame)	
	print("yaxis",yaxis)

	
	#计算各相关字段的偏离程度
	offset_result_start=offset_fu(xasis_fieldname_start,yaxis)
	print("字段英文名")
	print(xasis_fieldname_start)
	print("偏离程度")
	print(offset_result_start)
	print("字段相关系数值数组")
	print(regression_relation_start)
	

	#相关字段在正常范围内的不进行分析
	xasis_fieldname=[]
	offset_result=[]
	regression_relation=[]
	offsets_list=list(zip(xasis_fieldname_start,offset_result_start,regression_relation_start))
	print(offsets_list)
	print(offsets_list[0][1])

	for i in range(len(xasis_fieldname_start)):
		if offsets_list[i][1]==None:
			continue;
		elif abs(offsets_list[i][1])>0.2:
			xasis_fieldname.append(offsets_list[i][0])
			offset_result.append(offsets_list[i][1])
			regression_relation.append(offsets_list[i][2])			

	print("偏离程度大于20%的字段英文名")
	print(xasis_fieldname)
	print("偏离程度大于20%的字段偏离程度")
	print(offset_result)
	print("偏离程度大于20%的字段相关字段")
	print(regression_relation)		

	#正相关取偏高的/负相关取偏低的
	xasis_fieldname_result=[]#关联字段英文名字数组
	regression_relation_result=[]#字段相关系数值数组
	offset_result_final=[]#关联字段偏离程度值	
	if float(offset_value[0:-1])>=0:#读取质量字段偏离程度，由于隐藏域的偏离表示为百分比，例如12.6%。因此截取12.6来判断其正负
		for i in range(len(xasis_fieldname)):
			
			if float(regression_relation[i]) * float(offset_result[i]) >=0:
				xasis_fieldname_result.append(xasis_fieldname[i])
				regression_relation_result.append(regression_relation[i])
				offset_result_final.append(offset_result[i])
	else:
		for i in range(len(xasis_fieldname)):
			if float(regression_relation[i]) * float(offset_result[i]) <=0:
				xasis_fieldname_result.append(xasis_fieldname[i])
				regression_relation_result.append(regression_relation[i])
				offset_result_final.append(offset_result[i])
	
	print("最终相关字段名")
	print(xasis_fieldname_result)
	print("相关系数")
	print(regression_relation_result)
	print("偏离程度")
	print(offset_result_final)
	

	#取最大的前8个相关系数
	xasis_fieldname_result_max=xasis_fieldname_result[0:8]
	print('取最大的前8个相关字段')
	print(xasis_fieldname_result_max)

	#计算回归系数
	regression_coefficient=batchprocess.regression(field,xasis_fieldname_result_max,1)
	if regression_coefficient== False:
		return None,None,None,None
	coef=regression_coefficient[0]
	intercept=regression_coefficient[1]

	print('与前8个最终相关字段对应的回归系数')
	print(coef)
	print('截距')
	print(intercept)

	
	#取前3个权重最大的字段
	#按权重大小对8个字段进行排序
	coef_abs=[abs(float(n)) for n in coef]	
	print(coef_abs)
	print(type(coef_abs))
	M=dict(zip(coef,xasis_fieldname_result_max))
	print('组合前8个最终相关字段和回归系数')
	print(M)
	dicty_coef=dict(zip(coef_abs,xasis_fieldname_result_max))
	dicty_coef_order=sorted(dicty_coef.items(),reverse=True)
	print('组合前8个最终相关字段和回归系数按权重大小排序')
	print(dicty_coef_order)
	# coef_order_max=dicty_coef_order[0:3]
	#修改为追溯2个因素
	coef_order_max=dicty_coef_order[0:2]
	coef_field=[]
	coef_value=[]
	for n in range(len(coef_order_max)):
		coef_field.append(coef_order_max[n][1])
		coef_value.append(coef_order_max[n][0])
	print('前3个权重最大的字段和权重')	
	print(coef_order_max)
	print('前3个权重最大的字段')
	print(coef_field)
	print('前3个权重最大的字段权重')
	print(coef_value)
	coef_value_re=["%.4f"%n for n in list(coef_value)]

	#取前3个权重最大的字段实际值和偏离程度
	
	offsets_coef=value_offset(coef_field,prime_cost)
	print('前3个权重最大的字段偏离程度')
	print(offsets_coef)	
	
	

	#按操作顺序（表结构）进行排序

	#读取中文字段名
	ana_result={}
	ana_result=zhuanlu.PRO_BOF_HIS_ALLFIELDS
	En_to_Ch_result=[]
	for i in range(len(coef_field)):
		En_to_Ch_result.append(ana_result[coef_field[i]])
	print("中文字段")	
	print(En_to_Ch_result)

	#读取带标记的中文字段名（用于排序）
	ana_result_score={}
	ana_result_score=zhuanlu.PRO_BOF_HIS_ALLFIELDS_SCORE
	En_to_Ch_result_score=[]
	for i in range(len(coef_field)):
		En_to_Ch_result_score.append(ana_result_score[coef_field[i]])
	print("带标记的中文字段")	
	print(En_to_Ch_result_score)

	#中文字段与偏离程度合并，用于联动排序
	dicty=dict(zip(En_to_Ch_result,offsets_coef))
	print('组合字段中文名和偏离程度')
	print(dicty)

	#按操作排序字段名
	a=[]
	coef_field_score=[]
	for i in range(len(coef_field)):
		a.append(ana_result_score[coef_field[i]])
	L=sorted(a,key=by_score)
	for i in range(len(coef_field)):
		coef_field_score.append(L[i][0])
	print('按操作排序后中文名')
	print(coef_field_score)

    #按操作联动排序偏离程度
	offsets_coef_order=[]
	for i in range(len(coef_field)):
		offsets_coef_order.append(dicty[coef_field_score[i]])
	offset_value_single_cof=["%.2f%%"%(abs(float(n))*100) for n in list(offsets_coef_order)]
	print('按操作排序后字段偏离程度')
	print(offset_value_single_cof)


	#定性判断
	pos_float=[ float(n) for n in list(offsets_coef_order)]
	offset_result_nature=qualitative_offset(pos_float)
	print("定性判断")
	print(offset_result_nature)	
  


	return coef_field_score,offset_result_nature,offset_value_single_cof,coef_value_re

#取实际值并计算偏离程度
def value_offset(coef_field,prime_cost):
	#取相关字段实际值
	str_sql=''
	for i in range(len(coef_field)):
		 str_sql=str_sql+','+coef_field[i]
		# str_sql=str_sql+','+ 'nvl('+coef_field[i]+',0)as'+coef_field[i]
		#nvl(C,0) as C
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO" +str_sql+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where HEAT_NO='"+prime_cost+"'";
	print(sqlVO["sql"])
	scrapy_records_actual=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	yaxis=[]#各相关性字段实际值
	for i in range(len(coef_field)):
		value = scrapy_records_actual[0].get(coef_field[i],None)
		if value != None :
			scrapy_records_actual[0][coef_field[i]] = Decimal(value)
		else:
			scrapy_records_actual[0][coef_field[i]] = 0#将空值None暂时以0填充
		yaxis.append(value)
	frame=DataFrame(scrapy_records_actual)
	print("frame",frame)	
	print("yaxis",yaxis)	
	#计算各相关字段的偏离程度
	offset_result=offset_fu(coef_field,yaxis)
	return offset_result
def bound_judge(column,value):
	bound_lows = dict()
	bound_highs = dict()
	sqlVO = {"db_name": 'l2own'}
	sql = 'select DATA_ITEM_EN,NUMERICAL_LOWER_BOUND,NUMERICAL_UPPER_BOUND from QG_USER.PRO_BOF_HIS_ALLSTRUCTURE WHERE  IF_ANALYSE_TEMP = 1'
	sqlVO["sql"] = sql
	rs = models.BaseManage().direct_select_query_orignal_sqlVO(sqlVO)
	for row in rs:    
		bound_lows['%s' % row[0]] = '%s' % row[1]
		bound_highs['%s' % row[0]] = '%s' % row[2]
	#取到字段的上下限    
	bound_low = float(bound_lows.get(column,-999999999999))
	bound_high = float(bound_highs.get(column,999999999999))
	if  (value>= bound_low) & (value <= bound_high):
		value_tag=1;
	else:
		value_tag=0;
	return value_tag


