from pandas import DataFrame
import pandas as pd
import numpy as np
from decimal import *
from docx import Document
from docx.shared import Inches
from django.shortcuts import render
from . import models
from . import batchprocess
from . import chyulia
from django.shortcuts import render
from django.http import HttpResponse,HttpResponseRedirect,StreamingHttpResponse
import os
import json
import math



def violent_analyse(request):
	print("violent_ananlyse")
	document = Document()
	paragraph=document.add_paragraph()
	#炉次	
	for i in range(10):
		prime_cost=str(1634230+i);
		str_cause=violent_analyse_to(prime_cost,document,paragraph)

	contentVO={
		'title':'测试',
		'state':'success'		
	}				
	return HttpResponse(json.dumps(contentVO),content_type='application/json')	

def violent_analyse_to(prime_cost,document,paragraph):
	str_heatno='-------------------炉次号:'+prime_cost+'-------------------\n'      
	paragraph.add_run(str_heatno)

	# xasis_fieldname_en=['STEELWGT','LDG_STEELWGT','STEEL_SLAG']
	# xasis_fieldname_ch=['出钢量','转炉煤气','钢渣']

	field_classification=['raw','material','product','alloy']
	# for k in range(2,len(field_classification)):
	for k in range(0,4):
		classify='【'+field_classification[k]+'】\n'
		paragraph.add_run(classify)

		if field_classification[k]=='raw':
			#原料
			xasis_fieldname_ch=['铁水重量','生铁','废钢总和','大渣钢','自产废钢','重型废钢','中型废钢']
			xasis_fieldname_en=['MIRON_WGT','COLDPIGWGT','SCRAPWGT_COUNT','SCRAP_96053101','SCRAP_96052200','SCRAP_16010101','SCRAP_16020101']
			danwei=['Kg','Kg','Kg','Kg','Kg','Kg','Kg']
		elif field_classification[k]=='material':
			#物料
			xasis_fieldname_ch=['总吹氧消耗','氮气耗量','1#烧结矿','石灰石_40-70mm','萤石_FL80','增碳剂','低氮增碳剂','石灰','轻烧白云石']
			xasis_fieldname_en=['SUM_BO_CSM','N2CONSUME','L96020400','L12010302','L12010601','L12020201','L12020301','L96040100','L96040200']
			danwei=['NM3','NM3','Kg','Kg','Kg','Kg','Kg','Kg','Kg']
		elif field_classification[k]=='product':
			#产品
			xasis_fieldname_ch=['出钢量','转炉煤气','钢渣']
			xasis_fieldname_en=['STEELWGT','LDG_STEELWGT','STEEL_SLAG']
			danwei=['Kg','NM3','Kg']
		else:
			#合金
			xasis_fieldname_ch=['硅铁_Si72-80%、AL≤2%(粒度10-60mm)','微铝硅铁_Si 72-80%、AL≤0.1%、Ti≤0.1%','硅锰合金_Mn 65-72%、Si 17-20%','高硅硅锰_Mn ≥60%、Si ≥27%','中碳铬铁']
			xasis_fieldname_en=['L13010101','L13010301','L13020101','L13020201','L13040400']
			danwei=['Kg','Kg','Kg','Kg','Kg']

		sqlVO={}
		sqlVO["db_name"]="l2own"
		# sqlVO["sql"]="SELECT HEAT_NO,nvl(C,0) as C,nvl(SI,0) as SI ,nvl(MN,0) as MN,nvl(P,0) as P ,nvl(S,0) as S, nvl(Fe,0) as Fe, nvl(STEELWGT,0) as STEELWGT,nvl(FINAL_TEMP_VALUE,0)as FINAL_TEMP_VALUE FROM qg_user.PRO_BOF_HIS_ALLFIELDS where heat_no='"+prime_cost+"'";
		# sqlVO["sql"] = "SELECT nvl(" + ",0) ,nvl(".join(xaxis)  + ",0) from QG_USER.PRO_BOF_HIS_ALLFIELDS where heat_no= '"+prime_cost+"'"
		sqlVO["sql"] = "SELECT " + ",".join(xasis_fieldname_en) + " from QG_USER.PRO_BOF_HIS_ALLFIELDS where heat_no= '"+prime_cost+"'"
		scrapy_records_single=models.BaseManage().direct_select_query_sqlVO(sqlVO)
		print('单炉次数据库查询结果：',len(scrapy_records_single))
		print(scrapy_records_single)
		if len(scrapy_records_single) == 0:
			return None
		# xasis_fieldname_single=['C','SI','MN','P','S','Fe','STEELWGT','FINAL_TEMP_VALUE']
		# xaxis=['C含量','SI含量','MN含量','P含量','S含量','Fe含量','重量','温度']
		# xasis_fieldname_single=['C','SI']
		# xaxis=['C含量','SI含量']
		yaxis_single=[]

		for i in range(len(xasis_fieldname_en)):
			value = scrapy_records_single[0].get(xasis_fieldname_en[i],None)
			print(value)
			if value != None :
				scrapy_records_single[0][xasis_fieldname_en[i]] = float(value)
			else:
				value = 0
			yaxis_single.append(float(value))#由于在数据库查询时已经将空值设为了0，因此正常情况不会出现字段空值的情况

		# frame=DataFrame(scrapy_records_single)

		# yaxis_single=[frame.C[0],frame.SI[0],frame.MN[0],frame.P[0],frame.S[0],frame.FINAL_TEMP_VALUE[0]]

		print('单炉次成本分析字段')
		print(xasis_fieldname_en)
		print("单炉次成本分析字段实际值")
		print(yaxis_single)
		
		#计算单炉次质量分析字段偏离程度
		offset_result_single=chyulia.offset(xasis_fieldname_en,yaxis_single)
		offset_value_single=["%.2f%%"%(n*100) for n in list(offset_result_single)]
		print('单炉次成本分析字段偏离度')
		print(offset_value_single)

		#分析字段偏离程度定性判断
		qualitative_offset_result=chyulia.qualitative_offset(offset_result_single)
		print('单炉次成本分析字段偏离度定向分析')
		print(offset_value_single)
		#写入word

		for i in range(len(xasis_fieldname_en)):
			print("进行%s的追溯"%(xasis_fieldname_en[i]))
			xaxis_chinese=xasis_fieldname_ch[i];
			field=xasis_fieldname_en[i];
			single_value=yaxis_single[i];
			offset_value=offset_value_single[i];
			qualitative_offset_result_single=qualitative_offset_result[i];

			if single_value==0:#实际值为0，表示原来为空置或0值，追溯无意义
				str_des='本炉次'+prime_cost+'的'+xaxis_chinese+'字段实际值为空或0，追溯无意义！\n' 
				paragraph.add_run(str_des)	
				paragraph.add_run('\n')	
				continue
			En_to_Ch_result_score,offset_result_nature,offset_value_single_cof,regression_coefficient_result=analy_cof(prime_cost,field,single_value,offset_value);		
			if 	En_to_Ch_result_score==None:
				str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+'，但进行回归分析时相关字段无数据！'
				paragraph.add_run(str_des)	
				paragraph.add_run('\n')	
				continue
			if abs(float(offset_result_single[i]))<=0.1:
				str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+',偏离度为'+offset_value+'。\n'      
				paragraph.add_run(str_des)
			else:
				str_des='本炉次'+prime_cost+'的'+xaxis_chinese+qualitative_offset_result_single+',实际值为'+str(single_value)+danwei[i]+',偏离度为'+offset_value+'。通过数据相关性分析发现，导致该问题的原因是:\n'      
				paragraph.add_run(str_des)	
				for i in range(len(En_to_Ch_result_score)):
					str_cause=En_to_Ch_result_score[i]+offset_result_nature[i]+offset_value_single_cof[i]+',权重为'+str(regression_coefficient_result[i])+'\n'
					paragraph.add_run(str_cause) 
			paragraph.add_run('\n')
	document.add_page_break()
	document.save('e:/demo_chen.docx')
	return 	str_cause		

from . import zhuanlu
def  analy_cof(prime_cost,field,single_value,offset_value):
	#从数据库读取相关字段并按照相关系数绝对值由大到小排序
	print('进行%s炉次下的%s字段的追溯'%(prime_cost,field))
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT * FROM pro_bof_his_relation_cof where OUTPUTFIELD='"+field +"'order by abs(COF) desc"
	#print(sqlVO["sql"])
	scrapy_records_relation=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	#print(scrapy_records)
	length_result1=len(scrapy_records_relation)
	if length_result1==0:#即无相关性字段
		return None,None,None,None


	xasis_fieldname=[]#字段英文名字数组
	correlation_coefficient=[]#字段相关系数值数组
	str_sql=''
	for i in range(length_result1):
		xasis_fieldname.append(scrapy_records_relation[i].get('MIDDLEFIELD', None))
		correlation_coefficient.append(scrapy_records_relation[i].get('COF', None))
		str_sql=str_sql+','+scrapy_records_relation[i].get('MIDDLEFIELD', None)
	
	#取相关字段实际值
	sqlVO={}
	sqlVO["db_name"]="l2own"
	sqlVO["sql"]="SELECT HEAT_NO" +str_sql+" FROM qg_user.PRO_BOF_HIS_ALLFIELDS where HEAT_NO='"+prime_cost+"'";
	print(sqlVO["sql"])
	scrapy_records_actual=models.BaseManage().direct_select_query_sqlVO(sqlVO)
	yaxis=[]#各相关性字段实际值
	for i in range(length_result1):
		value = scrapy_records_actual[0].get(xasis_fieldname[i],None)
		if value != None :
			scrapy_records_actual[0][xasis_fieldname[i]] = Decimal(value)
		else:
			value = 0#将空值None暂时以0填充
		yaxis.append(value)
	# frame=DataFrame(scrapy_records_actual)
	# print("frame",frame)	
	print("yaxis",yaxis)

	offset_degree=chyulia.offset(xasis_fieldname,yaxis)
	print("字段名字：")
	print(xasis_fieldname)
	print("相关性：")
	print(correlation_coefficient)
	print("偏离程度")
	print(offset_degree)
	xasis_fieldname_result=[]#字段英文名字数组
	correlation_coefficient_result=[]#字段回归系数值数组
	offset_degree_result=[]#偏离程度值
	#即字段偏离高，则正相关应对应偏高，负相关应对应偏低

	j=0#标志位，表示当前筛选后的有效字段个数
	if float(offset_value[0:-1])>=0:
		for i in range(length_result1):
			if j>=8:#取筛选后相关性最大的八个因素字段
				break
			if offset_degree[i]==None or xasis_fieldname[i]=='NB' or yaxis[i]==0 :#由于数据清洗的问题，暂且将NB字段如此处理，因为NB字段的所有数据均相同，导致数据清洗时将所有数据都清除了
				continue
			if float(correlation_coefficient[i]) * float(offset_degree[i]) >=0:
				j+=1
				xasis_fieldname_result.append(xasis_fieldname[i])
				correlation_coefficient_result.append(correlation_coefficient[i])
				offset_degree_result.append(offset_degree[i])
	else:
		for i in range(length_result1):
			if j>=8:
				break
			if offset_degree[i]==None or xasis_fieldname[i]=='NB' or yaxis[i]==0:
				continue
			if float(correlation_coefficient[i]) * float(offset_degree[i]) <=0:
				j+=1
				xasis_fieldname_result.append(xasis_fieldname[i])
				correlation_coefficient_result.append(correlation_coefficient[i])
				offset_degree_result.append(offset_degree[i])
	
	#正负相关筛选后的相关字段个数
	print('正负相关筛选后的相关字段个数',len(xasis_fieldname_result))
	print("8个字段英文名字：")
	print(xasis_fieldname_result)

	if len(xasis_fieldname_result) == 0:
		return None,None,None,None

	#计算回归系数:regression_coefficient[0]表示各回归值，regression_coefficient[1]表示截距
	regression_coefficient=batchprocess.regression(field,xasis_fieldname_result,'null')
	if regression_coefficient== False:
		return None,None,None,None

	#查询转炉工序字段名中英文对照
	ana_result={}
	ana_result=zhuanlu.PRO_BOF_HIS_ALLFIELDS
	En_to_Ch_result=[]
	for i in range(len(xasis_fieldname_result)):
		En_to_Ch_result.append(ana_result[xasis_fieldname_result[i]])
	# print("字段中文名字")
	# print(En_to_Ch_result)

	print("8个字段英文名字：")
	print(xasis_fieldname_result)
	print("相关性系数")
	print(correlation_coefficient_result)
	print("回归系数")
	print(regression_coefficient[0])
	print("偏离程度")
	print(offset_degree_result)

	#zip压缩:中文名、英文名、偏离程度、回归系数;注：python3之后zip函数返回的是迭代值，需要list强转
	L_zip=list(zip(En_to_Ch_result,xasis_fieldname_result,offset_degree_result,regression_coefficient[0],correlation_coefficient_result))
	print('L_zip:',list(L_zip))
	#按照回归系数进行排序(从大到小)
	L_zip.sort(key=lambda x:x[3],reverse=True)
	print('按照回归系数排序后的字段：',L_zip)
	#取回归系数前三的字段
	L_zip=L_zip[0:3]
	print('取前三个字段时的实际字段个数：',len(L_zip))
	#解压缩
	L_unzip=list(zip(*L_zip))
	En_to_Ch_result_max=L_unzip[0]#中文名
	xasis_fieldname_result_max=L_unzip[1]#英文名
	offset_degree_result_max=L_unzip[2]#偏离程度
	regression_coefficient_max=L_unzip[3]#回归系数（权重）

	# '''
	#判断字段的顺序（对筛选后的三个字段再根据字段再表结构中的顺序进行排序）
	ana_result_score={}
	ana_result_score=zhuanlu.PRO_BOF_HIS_ALLFIELDS_SCORE
	En_to_Ch_result_score=[]
	for i in range(len(xasis_fieldname_result_max)):
		En_to_Ch_result_score.append(ana_result_score[xasis_fieldname_result_max[i]][1])
	# print('表结构中字段的顺序',En_to_Ch_result_score)

	#zip压缩：中文名、英文名、偏离程度、字段顺序编号
	L_zip = list(zip(En_to_Ch_result_max,xasis_fieldname_result_max,offset_degree_result_max,En_to_Ch_result_score))
	#按照字段顺序进行排序（从小到大）
	L_zip.sort(key=lambda x:x[3])
	print('追溯结果的三个字段',L_zip)
	#解压缩
	L_unzip=list(zip(*L_zip))

	En_to_Ch_result_max=L_unzip[0]#中文名
	xasis_fieldname_result_max=L_unzip[1]#英文名
	offset_degree_result_max=L_unzip[2]#偏离程度
	regression_coefficient_max=L_unzip[3]#回归系数（权重）
	# '''

	offset_degree_result_max_final=["%.2f%%"%(abs(float(n))*100) for n in list(offset_degree_result_max)]#将偏离程度值转化为保留两位小数的百分数
	#简单定性判断偏离程度（偏高、偏低）
	pos_float=[ float(n) for n in list(offset_degree_result_max)]
	offset_result_nature=simple_offset(pos_float)



	print('分析字段：',field)
	# print('分析字段偏离程度定性判断')
	# print(qualitative_offset_result[0])
	# print('分析字段偏离程度')
	# print(offset_value_single[0])
	# print('按操作排序后回归系数中文名')
	# print(En_to_Ch_result_score)
	# print("简单定性判断")
	# print(offset_result_nature)
	# print('按操作排序后字段偏离程度')
	# print(offset_value_single_cof)
	return En_to_Ch_result_max,offset_result_nature,offset_degree_result_max_final,regression_coefficient_max


#对相关性字段展示做简单定性分析	正数为偏高负数为偏低
def simple_offset(offset_result_final):
	#偏离程度定性标准，例如-10%~10%为正常，10%~20%为偏高，20%~40%为高，40%以上为数据异常/极端数据
	offset_result_nature=[]#相关字段定性分析
	for i in range(len(offset_result_final)):
		if float(offset_result_final[i]>0):
			offset_result_nature.append('偏高')
		else:
			offset_result_nature.append('偏低')	
	return  offset_result_nature



# #暴力求解
# if __name__ == '__main__':
# 	violent_analyse()